from docx import Document
from docx import shared
import json
import os
import sys
import re
from win32com.client import Dispatch,constants
from win32com.client import gencache
from excel_convert import c_reg
from docx.enum.style import WD_STYLE_TYPE

def jsonLoad(r_file):
    with open(r_file, 'r') as f:
        jsonData = json.load(f)
    return jsonData


def buildRegClass(jsonData):
    reg_list = []
    for reg in jsonData["Regs"]:
        # print(reg["name"])
        reg_list.append(c_reg(reg["name"], reg["addr"], reg["defaultValue"]))
        for field in reg["fields"]:
            # print(field["fieldName"])
            if field["width"] == 1:
                pos_str = str(field["fieldLsb"])
            else:
                msbPos = int(field["width"]) + int(field["fieldLsb"]) - 1
                pos_str = (str(msbPos) + ":" + str(field["fieldLsb"]))
            # if not re.search(r':', pos_str):
                # print("with : width")
            # print(type(pos_str))
            reg_list[-1].add_field(field["fieldName"],
                                   str(pos_str),
                                   field["type"],
                                   field["description"])

    return (jsonData["moduleName"], jsonData["addrWidth"],jsonData["baseAddr"], reg_list)


def outputDocx(moduleName, addrWidth,baseAddr, reg_list):
    document = Document()
    document.styles['Normal'].font.size=shared.Pt(10)
    table = document.add_table(2, 7)
    table.style ="Table Grid"
    proc_cells = table.rows[0].cells
    proc_cells[0].text = 'ModuleName'
    proc_cells[1].text = moduleName
    proc_cells[2].text = 'BaseAddr'
    proc_cells[3].text = str(hex(baseAddr))
    proc_cells[4].text = 'AddrWidth'
    proc_cells[5].text = str(addrWidth)
    proc_cells[0].paragraphs[0].runs[0].font.bold = True
    proc_cells[2].paragraphs[0].runs[0].font.bold = True
    proc_cells[4].paragraphs[0].runs[0].font.bold = True
    # for cell in proc_cells:
    #     if (len(cell.paragraphs[0].runs) >0 ):
    #         cell.paragraphs[0].runs[0].font.size = shared.Pt(10)
    proc_cells = table.rows[1].cells
    proc_cells[0].text = 'Name'
    proc_cells[1].text = 'OfstAddr'
    proc_cells[2].text = 'DefaultValue'
    proc_cells[3].text = 'FieldName'
    proc_cells[4].text = 'FieldWidth'
    proc_cells[5].text = 'Type'
    proc_cells[6].text = 'Description'
    for cell in proc_cells:
         cell.paragraphs[0].runs[0].font.bold = True
    #     #print(cell.paragraphs[0].runs[0].font.size)# = shared.Pt(24)
    #     #print(shared.Pt(24))
    #     cell.paragraphs[0].runs[0].font.size = shared.Pt(10)

    for reg in reg_list:
        table.add_row()
        dval_adj = reg.get_default_val()
        dval_adj = hex(dval_adj)
        dval_adj = re.sub("0x", "", dval_adj)
        dval_adj = dval_adj.rjust(4, '0')
        dval_adj = "32" + "'h" + dval_adj

        addr_adj = hex(reg.get_addr())
        addr_adj = re.sub("0x", "", addr_adj)
        addr_adj = addr_adj.rjust(4, '0')
        addr_adj = str(addrWidth) + "'h" + addr_adj

        proc_cells = table.rows[-1].cells
        proc_cells[0].text = reg.get_name()
        proc_cells[1].text = addr_adj
        proc_cells[2].text = dval_adj
        # proc_cells[5].merge(proc_cells[6])
        # proc_cells[4].merge(proc_cells[5])
        # proc_cells[3].merge(proc_cells[4])
        cur_field_idx = 31
        for field in reg.get_field_list():
            if(field["width"]+field["begin_pos"]-1 == cur_field_idx):
                if(cur_field_idx != 31):
                    table.add_row()
                proc_cells = table.rows[-1].cells
                proc_cells[3].text = field["name"]
                if field["width"] > 1:
                    proc_cells[4].text = str(field["width"]+field["begin_pos"]-1)+":"+str(field["begin_pos"])
                else:
                    proc_cells[4].text = str(field["begin_pos"])
                proc_cells[5].text = field["type"]
                proc_cells[6].text = field["description"]
                cur_field_idx -= field["width"]
            else:
                if(cur_field_idx != 31):
                    table.add_row()
                proc_cells = table.rows[-1].cells
                proc_cells[3].text = "Reserved"
                proc_cells[4].text = str(cur_field_idx) + ":" + str(field["width"]+field["begin_pos"])
                proc_cells[5].text = "RO"
                proc_cells[6].text = "保留"
                table.add_row()
                proc_cells = table.rows[-1].cells
                proc_cells[3].text = field["name"]
                if field["width"] > 1:
                    proc_cells[4].text = str(field["width"]+field["begin_pos"]-1)+":"+str(field["begin_pos"])
                else:
                    proc_cells[4].text = str(field["begin_pos"])
                proc_cells[5].text = field["type"]
                proc_cells[6].text = field["description"]
                cur_field_idx = field["begin_pos"]-1

    document.save(moduleName+'.docx')
def outputPDF(moduleName):
    # word = Dispatch("Word.Application")
    # word.Visible = 0
    # word.DisplayAlerts = 0
    # doc = word.Documents.Open(os.getcwd()+"\\"+moduleName+".docx")
    #
    # doc.SaveAS(os.getcwd()+"\\" + moduleName + ".pdf",FileFormat = 17)
    # doc.Close()
    # word.Quit()
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(os.getcwd()+"\\"+moduleName+".docx",ReadOnly=1)
    doc.ExportAsFixedFormat(os.getcwd()+"\\" + moduleName + ".pdf",
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)
if __name__ == "__main__":
    for file in sys.argv[1:len(sys.argv)-1]:
        if not os.path.exists(file):
            sys.exit("ERROR! file: %s do not exist!" % file)
        if not os.path.isfile(file):
            sys.exit("ERROR! file: %s is not a file!" % file)

        (moduleName, addrWidth,baseAddr, reg_list) = buildRegClass(jsonLoad(file))

        if (sys.argv[-1] == "json2docx"):
            outputDocx(moduleName, addrWidth, baseAddr, reg_list)
            outputPDF(moduleName)
        else:
            sys.exit("Error Function Input!")


